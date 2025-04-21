import os
import glob
from dotenv import load_dotenv

# Word ドキュメント読み込み用
from docx import Document as DocxDocument

# LlamaIndex の主要モジュール
from llama_index.core import Document, Settings, PromptHelper
from llama_index.embeddings.huggingface import HuggingFaceEmbedding

# 修正：GoogleGenAI -> GoogleGenerativeAI に変更
from langchain_google_genai import GoogleGenerativeAI

# .env ファイルから環境変数を読み込み（必要なら）
load_dotenv()

# Word ファイルが配置されるディレクトリと、インデックスの永続化先ディレクトリ
WORD_DIR = "./static/sample_FAQ_document"
INDEX_DB_DIR = "./static/vector_db_llamaindex"
os.makedirs(INDEX_DB_DIR, exist_ok=True)

# テキスト分割用パラメータ
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200  # チャンク間の重複数

def split_text(text: str, chunk_size: int = CHUNK_SIZE, chunk_overlap: int = CHUNK_OVERLAP) -> list[str]:
    """与えられたテキストを固定長チャンクに分割するシンプルな実装"""
    chunks = []
    start = 0
    text_length = len(text)
    while start < text_length:
        end = min(text_length, start + chunk_size)
        chunks.append(text[start:end])
        if end == text_length:
            break
        start = end - chunk_overlap
    return chunks

def process_docx(docx_path: str) -> list[Document]:
    """
    .docx を読み込み、段落ごとにテキストをチャンクに分割して
    Document オブジェクトのリストを生成する。
    各 Document には、元ファイル名と段落番号がメタデータとして付与される。
    """
    documents = []
    try:
        doc = DocxDocument(docx_path)
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            metadata = {
                "source": os.path.basename(docx_path),
                "paragraph": para_idx + 1
            }
            for chunk in split_text(text):
                documents.append(Document(text=chunk, extra_info=metadata))
        return documents
    except Exception as e:
        print(f"Error processing {docx_path}: {e}")
        return []

# LLM のインスタンスを生成し、Settings に設定
llm = GoogleGenerativeAI(model="gemini-2.0-flash")
Settings.llm = llm

# PromptHelper の初期化（max_tokens, chunk_size, chunk_overlap_ratio）
prompt_helper = PromptHelper(4096, CHUNK_SIZE, CHUNK_OVERLAP / CHUNK_SIZE)

# 埋め込みモデルを設定
embed_model = HuggingFaceEmbedding(model_name="intfloat/multilingual-e5-large")
Settings.embed_model = embed_model

def create_vector_indices():
    """
    Word (.docx) ファイルごとにベクトルインデックスを生成し、ディスクに永続化する関数
    """
    word_files = glob.glob(os.path.join(WORD_DIR, "*.docx"))
    if not word_files:
        raise FileNotFoundError(f"No Word files found in {WORD_DIR}")
    
    try:
        from llama_index import VectorStoreIndex
    except ImportError:
        from llama_index.core.indices.vector_store.base import VectorStoreIndex

    for word_file in word_files:
        name = os.path.splitext(os.path.basename(word_file))[0]
        subdir = os.path.join(INDEX_DB_DIR, name)
        os.makedirs(subdir, exist_ok=True)

        docs = process_docx(word_file)
        if not docs:
            print(f"No valid text in {word_file}. Skipping...")
            continue

        index = VectorStoreIndex.from_documents(
            docs,
            llm=llm,
            prompt_helper=prompt_helper,
            embed_model=embed_model
        )

        persist_dir = os.path.join(subdir, "persist")
        os.makedirs(persist_dir, exist_ok=True)
        index.storage_context.persist(persist_dir)
        print(f"Index saved for {name} to {persist_dir}")

if __name__ == "__main__":
    create_vector_indices()
